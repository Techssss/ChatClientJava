from __future__ import annotations

import argparse
import json
import math
import os
import textwrap
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "report_assets"
DIAGRAMS = ASSETS / "diagrams"
SCREENS = ASSETS / "screens"

BLUE = "2E74B5"
DARK_BLUE = "16324F"
INK = "243247"
MUTED = "667085"
LIGHT = "F4F6F9"
LIGHT_BLUE = "EAF2FB"
TEAL = "20B57E"
RED = "C0392B"
GOLD = "B7791F"
WHITE = "FFFFFF"


def font_path(bold=False):
    candidates = [
        Path("C:/Windows/Fonts/timesbd.ttf" if bold else "C:/Windows/Fonts/times.ttf"),
        Path("C:/Windows/Fonts/seguisb.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf"),
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return str(path)
    return None


def pil_font(size, bold=False):
    path = font_path(bold)
    return ImageFont.truetype(path, size=size) if path else ImageFont.load_default()


class Diagram:
    def __init__(self, width, height, title):
        self.w, self.h = width, height
        self.image = Image.new("RGB", (width, height), "white")
        self.d = ImageDraw.Draw(self.image)
        self.title_font = pil_font(42, True)
        self.h_font = pil_font(27, True)
        self.body_font = pil_font(23)
        self.small_font = pil_font(19)
        self.d.text((width / 2, 28), title, font=self.title_font, fill="#16324F", anchor="ma")

    def wrapped(self, text, width_px, font=None):
        font = font or self.body_font
        words = str(text).split()
        lines, current = [], ""
        for word in words:
            probe = word if not current else current + " " + word
            if self.d.textlength(probe, font=font) <= width_px:
                current = probe
            else:
                if current:
                    lines.append(current)
                current = word
        if current:
            lines.append(current)
        return lines or [""]

    def box(self, xy, title, lines=None, fill="#EAF2FB", outline="#2E74B5", radius=22,
            title_fill="#16324F", body_fill="#243247", center=False):
        x1, y1, x2, y2 = xy
        self.d.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=3)
        anchor = "ma" if center else "la"
        tx = (x1 + x2) / 2 if center else x1 + 22
        self.d.text((tx, y1 + 18), title, font=self.h_font, fill=title_fill, anchor=anchor)
        if lines:
            y = y1 + 58
            for line in lines:
                for part in self.wrapped(line, x2 - x1 - 38, self.small_font):
                    self.d.text((tx if center else x1 + 22, y), part, font=self.small_font,
                                fill=body_fill, anchor="ma" if center else "la")
                    y += 27

    def ellipse(self, xy, text, fill="#F4F6F9", outline="#2E74B5"):
        self.d.ellipse(xy, fill=fill, outline=outline, width=3)
        x1, y1, x2, y2 = xy
        lines = self.wrapped(text, x2 - x1 - 35, self.small_font)
        total = len(lines) * 25
        y = (y1 + y2) / 2 - total / 2 + 12
        for line in lines:
            self.d.text(((x1+x2)/2, y), line, font=self.small_font, fill="#243247", anchor="mm")
            y += 25

    def arrow(self, start, end, label=None, color="#5B6CFF", width=4, dashed=False):
        x1, y1 = start
        x2, y2 = end
        if dashed:
            steps = 18
            for i in range(0, steps, 2):
                a, b = i/steps, min((i+1)/steps, 1)
                self.d.line((x1+(x2-x1)*a, y1+(y2-y1)*a,
                             x1+(x2-x1)*b, y1+(y2-y1)*b), fill=color, width=width)
        else:
            self.d.line((x1, y1, x2, y2), fill=color, width=width)
        angle = math.atan2(y2-y1, x2-x1)
        size = 14
        pts = [(x2, y2),
               (x2-size*math.cos(angle-math.pi/6), y2-size*math.sin(angle-math.pi/6)),
               (x2-size*math.cos(angle+math.pi/6), y2-size*math.sin(angle+math.pi/6))]
        self.d.polygon(pts, fill=color)
        if label:
            mx, my = (x1+x2)/2, (y1+y2)/2
            bbox = self.d.textbbox((0, 0), label, font=self.small_font)
            pad = 7
            self.d.rounded_rectangle((mx-(bbox[2]-bbox[0])/2-pad, my-18,
                                      mx+(bbox[2]-bbox[0])/2+pad, my+14),
                                     radius=6, fill="white")
            self.d.text((mx, my-2), label, font=self.small_font, fill="#344054", anchor="mm")

    def actor(self, x, y, label):
        self.d.ellipse((x-24, y, x+24, y+48), outline="#16324F", width=4)
        self.d.line((x, y+48, x, y+132), fill="#16324F", width=4)
        self.d.line((x-48, y+78, x+48, y+78), fill="#16324F", width=4)
        self.d.line((x, y+132, x-42, y+188), fill="#16324F", width=4)
        self.d.line((x, y+132, x+42, y+188), fill="#16324F", width=4)
        self.d.text((x, y+216), label, font=self.h_font, fill="#16324F", anchor="ma")

    def save(self, name):
        DIAGRAMS.mkdir(parents=True, exist_ok=True)
        path = DIAGRAMS / name
        self.image.save(path, dpi=(180, 180))
        return path


def generate_use_case():
    g = Diagram(1800, 1120, "BIỂU ĐỒ USE CASE TỔNG QUÁT")
    g.actor(130, 420, "Người dùng")
    g.actor(1670, 420, "Quản trị server")
    boundary = (300, 110, 1500, 1050)
    g.d.rounded_rectangle(boundary, radius=30, outline="#98A2B3", width=4)
    g.d.text((330, 132), "Hệ thống Chatly", font=g.h_font, fill="#667085")
    cases = [
        (420, 210, 780, 320, "Đăng ký tài khoản"),
        (420, 350, 780, 460, "Đăng nhập"),
        (420, 490, 780, 600, "Xem người dùng online"),
        (420, 630, 780, 740, "Chat riêng tư mã hóa"),
        (420, 770, 780, 880, "Gửi ảnh, file, sticker"),
        (830, 210, 1190, 320, "Gửi ảnh giấu tin"),
        (830, 350, 1190, 460, "Speech-to-Text"),
        (830, 490, 1190, 600, "Gọi thoại"),
        (830, 630, 1190, 740, "Gọi video"),
        (830, 770, 1190, 880, "Tìm kiếm hội thoại"),
        (1210, 350, 1450, 470, "Bật / tắt server"),
        (1210, 540, 1450, 660, "Theo dõi nhật ký"),
    ]
    for x1, y1, x2, y2, text in cases:
        g.ellipse((x1, y1, x2, y2), text)
    for y in [265, 405, 545, 685, 825]:
        g.arrow((180, 510), (420, y), color="#667085", width=2)
    for y in [265, 405, 545, 685, 825]:
        g.arrow((180, 510), (830, y), color="#667085", width=2)
    g.arrow((1620, 510), (1450, 410), color="#667085", width=2)
    g.arrow((1620, 510), (1450, 600), color="#667085", width=2)
    g.d.text((900, 975), "Các use case nghiệp vụ chạy qua TCP port 5000 và SQLite phía server",
             font=g.small_font, fill="#667085", anchor="ma")
    return g.save("01-use-case.png")


def generate_architecture():
    g = Diagram(1800, 1080, "KIẾN TRÚC PHÂN LỚP VÀ THÀNH PHẦN")
    layers = [
        (100, 130, 1700, 300, "PRESENTATION", [
            (150, 180, 470, 270, "LoginUI", "Đăng nhập / đăng ký"),
            (540, 180, 860, 270, "ClientUI", "Chat và danh sách online"),
            (930, 180, 1250, 270, "Call UI", "Voice / video controls"),
            (1320, 180, 1640, 270, "ServerUI", "Console và trạng thái"),
        ], "#EEF2FF"),
        (100, 340, 1700, 520, "APPLICATION & MEDIA", [
            (150, 395, 470, 490, "ChatClient", "Socket client"),
            (540, 395, 860, 490, "AudioHandler", "Java Sound API"),
            (930, 395, 1250, 490, "VideoHandler", "Webcam capture"),
            (1320, 395, 1640, 490, "SttHandler", "Vosk offline STT"),
        ], "#ECFDF3"),
        (100, 560, 1700, 740, "SERVER & SECURITY", [
            (150, 615, 470, 710, "ChatServer", "Accept client / routing"),
            (540, 615, 860, 710, "ClientHandler", "Protocol dispatcher"),
            (930, 615, 1250, 710, "AES + RSA", "Hybrid encryption"),
            (1320, 615, 1640, 710, "DAO layer", "UserDAO / MessageDAO"),
        ], "#FFF7ED"),
        (100, 780, 1700, 980, "DATA & EXTERNAL RESOURCES", [
            (180, 845, 540, 945, "SQLite", "users / messages / message_keys"),
            (720, 845, 1080, 945, "Thiết bị", "Microphone / speaker / webcam"),
            (1260, 845, 1620, 945, "Vosk model", "Nhận dạng giọng nói offline"),
        ], "#F9FAFB"),
    ]
    for x1, y1, x2, y2, label, boxes, fill in layers:
        g.d.rounded_rectangle((x1, y1, x2, y2), radius=24, fill=fill, outline="#D0D5DD", width=3)
        g.d.text((x1+22, y1+16), label, font=g.small_font, fill="#667085")
        for bx1, by1, bx2, by2, title, sub in boxes:
            g.box((bx1, by1, bx2, by2), title, [sub], fill="white", outline="#98A2B3", radius=16, center=True)
    for y1, y2 in [(300, 340), (520, 560), (740, 780)]:
        for x in [310, 700, 1090, 1480]:
            g.arrow((x, y1), (x, y2), color="#5B6CFF", width=3)
    return g.save("02-architecture.png")


def generate_auth_activity():
    g = Diagram(1500, 1180, "BIỂU ĐỒ HOẠT ĐỘNG ĐĂNG NHẬP / ĐĂNG KÝ")
    cx = 750
    g.d.ellipse((cx-90, 100, cx+90, 160), fill="#16324F")
    g.d.text((cx, 130), "Bắt đầu", font=g.small_font, fill="white", anchor="mm")
    g.box((560, 210, 940, 300), "Nhập địa chỉ server", ["IP, port, username, password"], center=True)
    g.d.polygon([(cx, 350), (930, 440), (cx, 530), (570, 440)], fill="#FFF7ED", outline="#B7791F")
    g.d.text((cx, 440), "Chế độ?", font=g.h_font, fill="#7A5A00", anchor="mm")
    g.box((210, 590, 590, 700), "Create account", ["Kiểm tra xác nhận mật khẩu", "Gửi REGISTER"], fill="#ECFDF3", outline="#20B57E", center=True)
    g.box((910, 590, 1290, 700), "Sign in", ["Băm mật khẩu SHA-256", "Gửi LOGIN"], fill="#EEF2FF", outline="#5B6CFF", center=True)
    g.d.polygon([(cx, 760), (930, 850), (cx, 940), (570, 850)], fill="#F4F6F9", outline="#667085")
    g.d.text((cx, 850), "Server trả OK?", font=g.h_font, fill="#344054", anchor="mm")
    g.box((180, 990, 570, 1080), "Hiển thị lỗi", ["Trùng username / sai mật khẩu / không tồn tại"], fill="#FEF3F2", outline="#C0392B", center=True)
    g.box((930, 990, 1320, 1080), "Mở ClientUI", ["Kết nối socket mới và gửi CONNECT"], fill="#ECFDF3", outline="#20B57E", center=True)
    g.arrow((cx, 160), (cx, 210))
    g.arrow((cx, 300), (cx, 350))
    g.arrow((570, 440), (400, 590), "Đăng ký")
    g.arrow((930, 440), (1100, 590), "Đăng nhập")
    g.arrow((400, 700), (650, 790))
    g.arrow((1100, 700), (850, 790))
    g.arrow((570, 850), (375, 990), "Không")
    g.arrow((930, 850), (1125, 990), "Có")
    return g.save("03-auth-activity.png")


def sequence_base(title, participants, height=1250):
    g = Diagram(1900, height, title)
    margin = 150
    xs = [margin + i*(1600/(len(participants)-1)) for i in range(len(participants))]
    for x, name in zip(xs, participants):
        g.d.rounded_rectangle((x-125, 110, x+125, 185), radius=14, fill="#EAF2FB", outline="#2E74B5", width=3)
        g.d.text((x, 148), name, font=g.small_font, fill="#16324F", anchor="mm")
        g.d.line((x, 185, x, height-70), fill="#98A2B3", width=2)
    return g, xs


def seq_msg(g, xs, y, src, dst, label, response=False):
    color = "#20B57E" if response else "#5B6CFF"
    g.arrow((xs[src], y), (xs[dst], y), label, color=color, width=3, dashed=response)


def generate_auth_sequence():
    g, x = sequence_base("BIỂU ĐỒ TUẦN TỰ XÁC THỰC TÀI KHOẢN",
                         ["Người dùng", "LoginUI", "ClientHandler", "UserDAO", "SQLite"], 1300)
    steps = [
        (245, 0, 1, "Chọn Sign in / Create account", False),
        (330, 1, 1, "Kiểm tra dữ liệu và SHA-256", False),
        (420, 1, 2, "LOGIN hoặc REGISTER", False),
        (510, 2, 3, "findByUsername(username)", False),
        (600, 3, 4, "SELECT users", False),
        (680, 4, 3, "UserRecord / null", True),
        (765, 2, 3, "INSERT user + RSA keys (nếu đăng ký)", False),
        (850, 3, 4, "COMMIT", False),
        (940, 2, 1, "LOGIN_OK / REGISTER_OK + key", True),
        (1030, 1, 1, "Khởi tạo crypto context", False),
        (1115, 1, 2, "Socket mới: CONNECT", False),
        (1200, 2, 1, "USER_LIST", True),
    ]
    for args in steps:
        seq_msg(g, x, *args)
    return g.save("04-auth-sequence.png")


def generate_message_sequence():
    g, x = sequence_base("BIỂU ĐỒ TUẦN TỰ GỬI TIN NHẮN MÃ HÓA",
                         ["Sender UI", "ChatClient", "Server", "SQLite", "Receiver Client"], 1450)
    steps = [
        (235, 0, 1, "Nhập nội dung và chọn người nhận", False),
        (320, 1, 2, "KEY_REQUEST(receiver)", False),
        (405, 2, 3, "SELECT public_key", False),
        (490, 3, 2, "Public key", True),
        (575, 2, 1, "KEY_RESPONSE", True),
        (665, 0, 0, "Sinh AES-256 key + IV", False),
        (750, 0, 0, "AES-GCM(content); RSA-OAEP(AES key)", False),
        (840, 1, 2, "ENCRYPTED_TEXT", False),
        (930, 2, 3, "INSERT messages + message_keys", False),
        (1015, 2, 4, "Forward ciphertext", False),
        (1100, 4, 4, "RSA decrypt AES key", False),
        (1185, 4, 4, "AES-GCM decrypt content", False),
        (1270, 4, 4, "Hiển thị bubble", False),
    ]
    for args in steps:
        seq_msg(g, x, *args)
    g.d.text((950, 1385), "Lưu ý hiện trạng: receiverId cần được server ánh xạ từ username trước khi persist.",
             font=g.small_font, fill="#C0392B", anchor="mm")
    return g.save("05-message-sequence.png")


def generate_call_sequence():
    g, x = sequence_base("BIỂU ĐỒ TUẦN TỰ VOICE / VIDEO CALL",
                         ["Caller", "Server", "Callee", "Media Handler"], 1340)
    steps = [
        (250, 0, 1, "VOICE_CALL_REQ / VIDEO_CALL_REQ", False),
        (345, 1, 2, "Chuyển tiếp yêu cầu", False),
        (440, 2, 2, "Hiển thị hộp thoại Accept / Reject", False),
        (535, 2, 1, "CALL_RES(ACCEPT)", False),
        (630, 1, 0, "Chuyển tiếp ACCEPT", True),
        (725, 0, 3, "Mở microphone / webcam", False),
        (820, 2, 3, "Mở microphone / webcam", False),
        (915, 0, 1, "AUDIO_DATA / VIDEO_DATA", False),
        (1000, 1, 2, "Forward media frames", False),
        (1095, 2, 1, "Media frames chiều ngược lại", False),
        (1190, 0, 1, "CALL_END", False),
        (1270, 1, 2, "Dừng thiết bị và đóng UI", False),
    ]
    for args in steps:
        seq_msg(g, x, *args)
    return g.save("06-call-sequence.png")


def class_box(g, xy, name, attrs, methods, fill="#F8FAFC"):
    x1, y1, x2, y2 = xy
    g.d.rounded_rectangle(xy, radius=16, fill=fill, outline="#475467", width=3)
    g.d.rectangle((x1, y1, x2, y1+48), fill="#EAF2FB", outline="#475467", width=2)
    g.d.text(((x1+x2)/2, y1+24), name, font=g.h_font, fill="#16324F", anchor="mm")
    y = y1+62
    for a in attrs:
        g.d.text((x1+14, y), a, font=g.small_font, fill="#344054")
        y += 24
    g.d.line((x1, y+4, x2, y+4), fill="#98A2B3", width=2)
    y += 14
    for m in methods:
        g.d.text((x1+14, y), m, font=g.small_font, fill="#243247")
        y += 24


def generate_class_diagram():
    g = Diagram(2200, 1600, "BIỂU ĐỒ LỚP RÚT GỌN")
    boxes = {
        "LoginUI": (70, 130, 520, 380),
        "ClientUI": (620, 130, 1070, 460),
        "ChatClient": (1170, 130, 1600, 410),
        "Media": (1700, 130, 2130, 460),
        "ServerUI": (70, 590, 500, 810),
        "ChatServer": (610, 570, 1050, 850),
        "ClientHandler": (1160, 550, 1610, 900),
        "Message": (1710, 590, 2140, 900),
        "Crypto": (150, 1050, 650, 1370),
        "DAO": (850, 1050, 1350, 1390),
        "Database": (1570, 1080, 2070, 1370),
    }
    class_box(g, boxes["LoginUI"], "LoginUI", ["- authMode", "- username/password"],
              ["+ authenticate()", "+ authenticateViaServer()", "+ updateModeUi()"])
    class_box(g, boxes["ClientUI"], "ClientUI", ["- selectedUser", "- crypto context", "- chatPanels"],
              ["+ handleMessage()", "- sendText()", "- buildAndSendEncryptedMsg()", "- addMessageBubble()"])
    class_box(g, boxes["ChatClient"], "ChatClient", ["- Socket", "- Object streams", "- username"],
              ["+ connect()", "+ sendMessage()", "+ disconnect()", "+ run()"])
    class_box(g, boxes["Media"], "Audio / Video / STT", ["AudioHandler", "VideoHandler", "SttHandler"],
              ["+ startCall()", "+ playAudio()/displayVideo()", "+ startRecording()"])
    class_box(g, boxes["ServerUI"], "ServerUI", ["- ChatServer", "- status/log"],
              ["+ toggleServer()", "+ log()"])
    class_box(g, boxes["ChatServer"], "ChatServer", ["- clients: Map", "- ServerSocket"],
              ["+ addClient()", "+ removeClient()", "+ sendToUser()", "+ broadcastUserList()"])
    class_box(g, boxes["ClientHandler"], "ClientHandler", ["- UserDAO", "- MessageDAO", "- username"],
              ["+ run()", "- handleAuthentication()", "- handleMessage()", "- persistEncryptedMessage()"])
    class_box(g, boxes["Message"], "Message", ["type, sender, receiver", "content, fileData", "ciphertext, AES keys, IV"],
              ["Serializable", "getters / setters"])
    class_box(g, boxes["Crypto"], "Crypto utilities", ["AESUtil", "RSAUtil", "KeyManager"],
              ["AES-256-GCM", "RSA-2048 OAEP", "Base64 serialization"], fill="#ECFDF3")
    class_box(g, boxes["DAO"], "DAO layer", ["UserDAO", "MessageDAO", "CryptoMessageService"],
              ["create/find user", "save/query ciphertext", "decrypt history"], fill="#FFF7ED")
    class_box(g, boxes["Database"], "DatabaseManager", ["Singleton Connection", "DB_PATH=user.home"],
              ["+ getInstance()", "- initSchema()"], fill="#F4F6F9")
    links = [
        ("LoginUI", "ClientUI", "tạo"), ("ClientUI", "ChatClient", "sử dụng"),
        ("ClientUI", "Media", "điều phối"), ("ChatClient", "Message", "truyền"),
        ("ServerUI", "ChatServer", "khởi động"), ("ChatServer", "ClientHandler", "tạo nhiều"),
        ("ClientHandler", "Message", "xử lý"), ("ClientHandler", "DAO", "gọi"),
        ("ClientHandler", "Crypto", "tạo key"), ("DAO", "Database", "kết nối"),
    ]
    for a, b, label in links:
        A, B = boxes[a], boxes[b]
        start = ((A[0]+A[2])/2, A[3]) if B[1] > A[3] else (A[2], (A[1]+A[3])/2)
        end = ((B[0]+B[2])/2, B[1]) if B[1] > A[3] else (B[0], (B[1]+B[3])/2)
        g.arrow(start, end, label, color="#667085", width=2)
    return g.save("07-class-diagram.png")


def erd_table(g, xy, name, rows):
    x1, y1, x2, y2 = xy
    g.d.rounded_rectangle(xy, radius=14, fill="white", outline="#2E74B5", width=3)
    g.d.rectangle((x1, y1, x2, y1+55), fill="#2E74B5")
    g.d.text(((x1+x2)/2, y1+28), name, font=g.h_font, fill="white", anchor="mm")
    y = y1+68
    for key, field in rows:
        color = "#B7791F" if key == "PK" else ("#C0392B" if key == "FK" else "#667085")
        g.d.text((x1+18, y), key, font=g.small_font, fill=color)
        g.d.text((x1+78, y), field, font=g.small_font, fill="#243247")
        y += 34


def generate_erd():
    g = Diagram(1800, 1080, "BIỂU ĐỒ THỰC THỂ - QUAN HỆ (ERD)")
    erd_table(g, (90, 180, 620, 650), "users", [
        ("PK", "id INTEGER"), ("", "username TEXT UNIQUE"), ("", "password_hash TEXT"),
        ("", "public_key TEXT"), ("", "private_key TEXT"), ("", "created_at DATETIME")])
    erd_table(g, (720, 180, 1250, 610), "messages", [
        ("PK", "id INTEGER"), ("FK", "sender_id -> users.id"), ("FK", "receiver_id -> users.id"),
        ("", "encrypted_content TEXT"), ("", "iv TEXT"), ("", "created_at DATETIME")])
    erd_table(g, (1180, 700, 1710, 1010), "message_keys", [
        ("PK", "id INTEGER"), ("FK", "message_id -> messages.id"),
        ("FK", "user_id -> users.id"), ("", "encrypted_aes_key TEXT")])
    g.arrow((620, 310), (720, 310), "1 - N sender", color="#5B6CFF")
    g.arrow((620, 480), (720, 480), "1 - N receiver", color="#5B6CFF")
    g.arrow((1110, 610), (1350, 700), "1 - N", color="#20B57E")
    g.arrow((620, 590), (1180, 850), "1 - N user", color="#B7791F")
    g.d.text((900, 920), "Mỗi message có hai encrypted AES key: một cho sender và một cho receiver.",
             font=g.small_font, fill="#667085", anchor="mm")
    return g.save("08-erd.png")


def generate_deployment():
    g = Diagram(1800, 1050, "BIỂU ĐỒ TRIỂN KHAI HAI MÁY TRONG MẠNG LAN")
    g.box((80, 150, 820, 900), "Máy A - Server", ["Windows + JDK 17/21", "ServerUI / ChatServer", "TCP port 5000"],
          fill="#EEF2FF", outline="#5B6CFF")
    g.box((980, 150, 1720, 900), "Máy B - Client", ["Windows + JDK", "LoginUI / ClientUI", "Nhập IPv4 Wi-Fi của Máy A"],
          fill="#ECFDF3", outline="#20B57E")
    g.box((160, 390, 740, 540), "SQLite", ["C:/Users/<user>/chat_server.db", "Chỉ server truy cập"], fill="white", outline="#667085", center=True)
    g.box((160, 620, 740, 790), "Thiết bị tùy chọn", ["Microphone, speaker, webcam", "Vosk model"], fill="white", outline="#667085", center=True)
    g.box((1060, 390, 1640, 540), "Client runtime", ["Socket TCP + Object streams", "UI Swing / FlatLaf"], fill="white", outline="#667085", center=True)
    g.box((1060, 620, 1640, 790), "Thiết bị tùy chọn", ["Microphone, speaker, webcam", "Vosk model"], fill="white", outline="#667085", center=True)
    g.arrow((820, 500), (980, 500), "Wi-Fi LAN / TCP 5000", color="#5B6CFF", width=5)
    g.arrow((980, 570), (820, 570), "Message, file, audio, video", color="#20B57E", width=4)
    g.d.text((900, 965), "Windows Firewall trên Máy A phải cho phép Java hoặc inbound TCP 5000.",
             font=g.h_font, fill="#C0392B", anchor="mm")
    return g.save("09-deployment.png")


def generate_call_state():
    g = Diagram(1700, 950, "BIỂU ĐỒ TRẠNG THÁI CUỘC GỌI")
    states = {
        "Idle": (100, 380, 330, 520),
        "Calling": (430, 180, 720, 330),
        "Ringing": (430, 600, 720, 750),
        "Connected": (850, 380, 1180, 540),
        "MediaOff": (1280, 160, 1600, 330),
        "Ended": (1300, 620, 1580, 770),
    }
    labels = {
        "Idle": ("IDLE", ["Chưa có cuộc gọi"]),
        "Calling": ("CALLING", ["Đã gửi CALL_REQ"]),
        "Ringing": ("RINGING", ["Chờ Accept / Reject"]),
        "Connected": ("CONNECTED", ["Streaming media", "Timer hoạt động"]),
        "MediaOff": ("MUTED / CAMERA OFF", ["Giữ kết nối", "Tạm dừng track"]),
        "Ended": ("ENDED", ["Đóng thiết bị", "Giải phóng UI"]),
    }
    for key, xy in states.items():
        title, lines = labels[key]
        fill = "#ECFDF3" if key == "Connected" else ("#FEF3F2" if key == "Ended" else "#EEF2FF")
        outline = "#20B57E" if key == "Connected" else ("#C0392B" if key == "Ended" else "#5B6CFF")
        g.box(xy, title, lines, fill=fill, outline=outline, center=True)
    def center(k):
        x1, y1, x2, y2 = states[k]
        return ((x1+x2)/2, (y1+y2)/2)
    for a, b, label in [("Idle","Calling","Bấm gọi"), ("Calling","Ringing","Forward request"),
                        ("Ringing","Connected","Accept"), ("Calling","Ended","Reject / timeout"),
                        ("Connected","MediaOff","Mute / camera off"), ("MediaOff","Connected","Bật lại"),
                        ("Connected","Ended","End / mất kết nối"), ("Ended","Idle","Đóng cửa sổ")]:
        g.arrow(center(a), center(b), label, color="#667085", width=3)
    return g.save("10-call-state.png")


def generate_all_diagrams():
    return [
        generate_use_case(), generate_architecture(), generate_auth_activity(),
        generate_auth_sequence(), generate_message_sequence(), generate_call_sequence(),
        generate_class_diagram(), generate_erd(), generate_deployment(), generate_call_state(),
    ]


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def set_run_font(run, name="Times New Roman", size=None, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    tr_pr.append(element)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent=120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa)-1)]
            set_cell_width(cell, width)
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, color=BLUE, size=12, side="left"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), "8")
    border.set(qn("w:color"), color)
    p_bdr.append(border)


def set_paragraph_shading(paragraph, fill=LIGHT):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    # Named override: Vietnamese academic report geometry and typography.
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.5

    heading_tokens = {
        "Heading 1": (16, BLUE, 14, 7),
        "Heading 2": (14, BLUE, 11, 5),
        "Heading 3": (13, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(13)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.5

    caption = styles["Caption"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    caption.font.size = Pt(11)
    caption.font.italic = True
    caption.font.color.rgb = rgb(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.keep_together = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("BÁO CÁO ĐỒ ÁN  |  CHATLY")
    set_run_font(hr, size=9, color=MUTED, bold=True)
    first_header = section.first_page_header
    first_header.paragraphs[0].text = ""

    footer = section.footer
    ft = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(ft, [6300, 2772], indent=0)
    ft.style = None
    left = ft.cell(0, 0).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(left.add_run("ChatClientJava • Commit 1ded34a"), size=9, color=MUTED)
    right = ft.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(right.add_run("Trang "), size=9, color=MUTED)
    add_page_field(right)
    section.first_page_footer.paragraphs[0].text = ""

    doc.core_properties.title = "Báo cáo đồ án ChatClientJava - Chatly"
    doc.core_properties.subject = "Ứng dụng chat client-server bằng Java Swing"
    doc.core_properties.author = "Sinh viên thực hiện"
    doc.core_properties.keywords = "Java, Swing, Socket, SQLite, AES-GCM, RSA, Vosk, Webcam"


def add_para(doc, text="", bold_prefix=None, align=None, italic=False, color=INK,
             size=13, after=6, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.alignment = align or WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.keep_together = keep
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, name="Times New Roman", size=size, color=color, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, name="Times New Roman", size=size, color=color, italic=italic)
    else:
        r = p.add_run(text)
        set_run_font(r, name="Times New Roman", size=size, color=color, italic=italic)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    r = p.add_run(text)
    set_run_font(r, name="Times New Roman", size=13, color=INK)
    return p


def add_number(doc, text, level=0):
    p = doc.add_paragraph(style="List Number" if level == 0 else "List Number 2")
    r = p.add_run(text)
    set_run_font(r, name="Times New Roman", size=13, color=INK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    return p


def add_callout(doc, label, text, fill=LIGHT_BLUE, accent=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.2
    set_paragraph_shading(p, fill)
    set_paragraph_border(p, accent, 18, "left")
    set_run_font(p.add_run(label + ": "), name="Times New Roman", size=12, color=accent, bold=True)
    set_run_font(p.add_run(text), name="Times New Roman", size=12, color=INK)
    return p


def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.keep_together = True
    set_paragraph_shading(p, "F2F4F7")
    set_paragraph_border(p, "98A2B3", 10, "left")
    r = p.add_run(code)
    set_run_font(r, name="Consolas", size=8.5, color="273444")
    return p


def add_table(doc, headers, rows, widths, caption=None):
    if caption:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        set_run_font(p.add_run(caption), size=9, color=MUTED, italic=True)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, text in enumerate(headers):
        cell = header.cells[i]
        shade_cell(cell, LIGHT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(text), size=9.5, color=DARK_BLUE, bold=True)
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            cell = row.cells[i]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(str(value)) < 18 and i != 1 else WD_ALIGN_PARAGRAPH.LEFT
            set_run_font(p.add_run(str(value)), size=9.2, color=INK)
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(3)
    return table


def add_figure(doc, path, caption, width=5.65):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", caption)
    shape._inline.docPr.set("title", caption)
    cp = doc.add_paragraph(caption, style="Caption")
    return cp


def page_break(doc):
    doc.add_page_break()


TOC_ENTRIES = [
    "TÓM TẮT ĐỒ ÁN",
    "DANH MỤC HÌNH VÀ BẢNG",
    "CHƯƠNG 1. TỔNG QUAN ĐỀ TÀI",
    "CHƯƠNG 2. KHẢO SÁT VÀ ĐẶC TẢ YÊU CẦU",
    "CHƯƠNG 3. PHÂN TÍCH HỆ THỐNG",
    "CHƯƠNG 4. THIẾT KẾ HỆ THỐNG",
    "CHƯƠNG 5. CÀI ĐẶT VÀ XÂY DỰNG CHỨC NĂNG",
    "CHƯƠNG 6. KIỂM THỬ VÀ ĐÁNH GIÁ",
    "CHƯƠNG 7. TRIỂN KHAI VÀ HƯỚNG DẪN SỬ DỤNG",
    "CHƯƠNG 8. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN",
    "TÀI LIỆU THAM KHẢO",
    "PHỤ LỤC",
]


def add_cover(doc):
    for _ in range(2):
        add_para(doc, "", after=0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("TRƯỜNG [ĐIỀN TÊN TRƯỜNG]"), size=14, color=DARK_BLUE, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("KHOA / BỘ MÔN CÔNG NGHỆ THÔNG TIN"), size=14, color=DARK_BLUE, bold=True)
    for _ in range(3):
        add_para(doc, "", after=0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("BÁO CÁO ĐỒ ÁN"), size=20, color=GOLD, bold=True)
    p.paragraph_format.space_after = Pt(14)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_run_font(p.add_run("XÂY DỰNG ỨNG DỤNG CHAT CLIENT - SERVER"), size=24, color=DARK_BLUE, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_run_font(p.add_run("BẰNG JAVA SWING"), size=24, color=DARK_BLUE, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(34)
    set_run_font(p.add_run("Chatly / ChatClientJava"), size=17, color=BLUE, bold=True)

    meta = [
        "Sinh viên thực hiện: [Điền họ và tên]",
        "Mã số sinh viên: [Điền MSSV]",
        "Lớp: [Điền lớp]",
        "Giảng viên hướng dẫn: [Điền họ tên giảng viên]",
    ]
    for line in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(5)
        set_run_font(p.add_run(line), size=13, color=INK)
    for _ in range(3):
        add_para(doc, "", after=0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("Tháng 06 năm 2026"), size=13, color=MUTED, bold=True)
    page_break(doc)


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)


def add_toc(doc, toc_pages):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    set_run_font(p.add_run("MỤC LỤC"), size=18, color=DARK_BLUE, bold=True)
    table = doc.add_table(rows=0, cols=2)
    remove_table_borders(table)
    for entry in TOC_ENTRIES:
        row = table.add_row()
        left, right = row.cells
        p1 = left.paragraphs[0]
        p1.paragraph_format.space_after = Pt(3)
        set_run_font(p1.add_run(entry), size=12.5, color=INK,
                     bold=entry.startswith("CHƯƠNG"))
        p2 = right.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_run_font(p2.add_run(str(toc_pages.get(entry, "…"))), size=12.5, color=MUTED)
    set_table_geometry(table, [7900, 1172], indent=0)
    add_callout(doc, "Ghi chú", "Mục lục được dựng tĩnh từ bản render cuối để số trang khớp với tài liệu giao nộp.",
                fill="F8FAFC", accent=MUTED)
    page_break(doc)


def add_abstract(doc):
    add_heading(doc, "TÓM TẮT ĐỒ ÁN", 1)
    add_para(doc,
        "Đồ án xây dựng Chatly, một ứng dụng nhắn tin thời gian thực theo mô hình client - server bằng Java Swing. "
        "Hệ thống cho phép người dùng đăng ký, đăng nhập, xem danh sách tài khoản đang trực tuyến, chat riêng tư, "
        "gửi ảnh và tệp, gửi ảnh giấu tin, chuyển giọng nói thành văn bản, gọi thoại và gọi video trong mạng LAN. "
        "Server chịu trách nhiệm tiếp nhận kết nối TCP, định tuyến thông điệp, quản lý người dùng trực tuyến và truy cập "
        "cơ sở dữ liệu SQLite; client đảm nhiệm giao diện, thiết bị media và mã hóa nội dung trước khi truyền.")
    add_para(doc,
        "Đối với tin nhắn văn bản, hệ thống áp dụng mô hình mã hóa lai: nội dung được bảo vệ bằng AES-256-GCM, "
        "khóa AES được mã hóa bằng RSA-2048 OAEP cho người gửi và người nhận. Tài khoản, khóa và cấu trúc lưu ciphertext "
        "được thiết kế trong SQLite. Giao diện sử dụng FlatLaf, SVG vector và bố cục hiện đại; tính năng media sử dụng "
        "Java Sound API, Webcam Capture và mô hình Vosk offline.")
    add_para(doc,
        "Kết quả kiểm thử cho thấy dự án biên dịch thành công bằng Maven, luồng đăng ký và đăng nhập phân biệt rõ bốn "
        "trường hợp: đăng ký mới, đăng nhập đúng, đăng nhập sai mật khẩu và đăng ký trùng. Sản phẩm đáp ứng tốt mục tiêu "
        "minh họa socket đa luồng, giao diện desktop, mã hóa lai và tích hợp thiết bị. Báo cáo cũng chỉ ra các giới hạn "
        "cần hoàn thiện trước khi triển khai thực tế, gồm tải lại lịch sử chat, ánh xạ receiverId khi lưu tin, bảo vệ khóa "
        "riêng, kiểm soát kích thước dữ liệu và thay Java serialization bằng protocol an toàn hơn.")
    add_callout(doc, "Từ khóa", "Java Swing; TCP Socket; SQLite; AES-GCM; RSA-OAEP; Vosk; Webcam; Client-Server.")


def add_lists(doc):
    add_heading(doc, "DANH MỤC HÌNH VÀ BẢNG", 1)
    figures = [
        "Hình 2.1. Biểu đồ use case tổng quát",
        "Hình 3.1. Kiến trúc phân lớp và thành phần",
        "Hình 3.2. Hoạt động đăng nhập / đăng ký",
        "Hình 3.3. Tuần tự xác thực tài khoản",
        "Hình 3.4. Tuần tự gửi tin nhắn mã hóa",
        "Hình 3.5. Tuần tự voice / video call",
        "Hình 4.1. Biểu đồ lớp rút gọn",
        "Hình 4.2. Biểu đồ thực thể - quan hệ",
        "Hình 4.3. Biểu đồ trạng thái cuộc gọi",
        "Hình 5.1 - 5.4. Giao diện đăng ký, chat, voice call và video call",
        "Hình 7.1. Biểu đồ triển khai hai máy trong mạng LAN",
    ]
    for item in figures:
        add_bullet(doc, item)
    add_para(doc, "Các bảng chính gồm: bảng yêu cầu chức năng, ma trận use case, bảng công nghệ, mô tả lớp, "
             "lược đồ dữ liệu, message type, ca kiểm thử, rủi ro và kế hoạch phát triển.", after=4)


def chapter_1(doc):
    add_heading(doc, "CHƯƠNG 1. TỔNG QUAN ĐỀ TÀI", 1)
    add_heading(doc, "1.1. Lý do chọn đề tài", 2)
    add_para(doc,
        "Nhắn tin thời gian thực là bài toán điển hình để kết hợp nhiều kiến thức nền tảng của ngành Công nghệ thông tin: "
        "lập trình hướng đối tượng, lập trình mạng, đa luồng, cơ sở dữ liệu, bảo mật, xử lý âm thanh - hình ảnh và thiết kế "
        "giao diện. Khác với một ứng dụng CRUD thông thường, hệ thống chat phải duy trì kết nối lâu dài, xử lý nhiều client "
        "đồng thời và bảo đảm dữ liệu được chuyển đến đúng người nhận với độ trễ thấp.")
    add_para(doc,
        "Java phù hợp cho đề tài nhờ tính đa nền tảng, hệ sinh thái thư viện phong phú và API socket, thread, sound, image "
        "được tích hợp tốt. Swing tuy là công nghệ desktop truyền thống nhưng vẫn hữu ích trong môi trường học thuật vì giúp "
        "người học quan sát trực tiếp mô hình event-driven và tổ chức UI mà không phụ thuộc trình duyệt hoặc web server.")
    add_heading(doc, "1.2. Mục tiêu", 2)
    for item in [
        "Xây dựng một server TCP đa luồng có khả năng tiếp nhận và quản lý nhiều client.",
        "Tạo client desktop có giao diện đăng ký, đăng nhập, chat và gọi media dễ sử dụng.",
        "Thiết kế cơ sở dữ liệu SQLite cho tài khoản và dữ liệu tin nhắn mã hóa.",
        "Áp dụng AES-GCM và RSA-OAEP để minh họa mô hình mã hóa lai.",
        "Tích hợp gửi tệp, ảnh, steganography, voice call, video call và speech-to-text.",
        "Triển khai và kiểm thử trong mô hình một server - nhiều client trên cùng máy hoặc trong mạng LAN.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "1.3. Phạm vi", 2)
    add_para(doc,
        "Phạm vi hiện tại tập trung vào chat riêng tư giữa các tài khoản đang online. Server chạy tại một máy, mặc định "
        "lắng nghe cổng 5000; các client trong cùng mạng nhập IPv4 Wi-Fi của máy server để kết nối. Hệ thống chưa triển khai "
        "group chat, đồng bộ cloud, push notification, NAT traversal, web client hoặc mobile client.")
    add_heading(doc, "1.4. Kết quả đạt được", 2)
    add_para(doc,
        "Phiên bản hiện tại có đầy đủ source Maven, giao diện SVG hiện đại, account flow tách Sign in / Create account, "
        "console server, chat riêng tư, gửi media và cuộc gọi. Dự án được build thành công với Java target 17; smoke test "
        "xác thực trả đúng REGISTER_OK, LOGIN_OK, LOGIN_FAIL và REGISTER_FAIL. Source đã được quản lý trên nhánh main của "
        "repository ChatClientJava.")


def chapter_2(doc):
    add_heading(doc, "CHƯƠNG 2. KHẢO SÁT VÀ ĐẶC TẢ YÊU CẦU", 1)
    add_heading(doc, "2.1. Tác nhân và nhu cầu", 2)
    add_para(doc,
        "Hệ thống có hai tác nhân chính. Người dùng cuối thao tác trên client để quản lý tài khoản, chọn người đang online, "
        "trao đổi nội dung và điều khiển cuộc gọi. Quản trị server khởi động hoặc dừng server, quan sát nhật ký kết nối và "
        "bảo đảm cổng mạng cùng cơ sở dữ liệu hoạt động.")
    add_figure(doc, DIAGRAMS / "01-use-case.png", "Hình 2.1. Biểu đồ use case tổng quát của hệ thống Chatly")
    add_heading(doc, "2.2. Yêu cầu chức năng", 2)
    requirements = [
        ("FR-01", "Đăng ký", "Tạo username duy nhất; mật khẩu tối thiểu 6 ký tự và xác nhận khớp.", "Bắt buộc"),
        ("FR-02", "Đăng nhập", "Xác thực username và password hash; báo lỗi rõ ràng.", "Bắt buộc"),
        ("FR-03", "Online list", "Server phát danh sách tài khoản đang kết nối.", "Bắt buộc"),
        ("FR-04", "Chat mã hóa", "Gửi tin riêng tư với AES-GCM và RSA-OAEP.", "Bắt buộc"),
        ("FR-05", "Ảnh / file", "Chọn file từ máy và chuyển tiếp đến người nhận.", "Cao"),
        ("FR-06", "Steganography", "Giấu và đọc văn bản trong ảnh PNG/BMP bằng mật khẩu.", "Trung bình"),
        ("FR-07", "Voice call", "Yêu cầu, chấp nhận, truyền audio, mute/speaker/end.", "Cao"),
        ("FR-08", "Video call", "Truyền video và audio, self-preview, camera/mic/speaker/end.", "Cao"),
        ("FR-09", "Speech-to-Text", "Nhận dạng giọng nói offline bằng Vosk.", "Trung bình"),
        ("FR-10", "Server console", "Start/stop, hiển thị trạng thái và log kết nối.", "Bắt buộc"),
    ]
    add_table(doc, ["Mã", "Chức năng", "Mô tả", "Ưu tiên"], requirements,
              [900, 1700, 5000, 1472], "Bảng 2.1. Danh sách yêu cầu chức năng")
    add_heading(doc, "2.3. Yêu cầu phi chức năng", 2)
    for item in [
        "Khả dụng: UI phản hồi trên Event Dispatch Thread; tác vụ mạng và mã hóa chạy nền.",
        "Hiệu năng: video giảm kích thước và gửi khoảng 10 khung hình/giây; audio dùng buffer 1024 byte.",
        "An toàn dữ liệu: text sử dụng authenticated encryption; password không lưu dạng rõ.",
        "Khả chuyển: source target Java 17; chạy trên Windows với Maven và JDK phù hợp.",
        "Bảo trì: tách package client, server, common, crypto và db; có UITheme dùng chung.",
        "Khả quan sát: server console ghi nhận start/stop, đăng ký, đăng nhập và mất kết nối.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "2.4. Ràng buộc", 2)
    add_para(doc,
        "Hai máy phải nhìn thấy nhau trong cùng mạng LAN hoặc có định tuyến phù hợp; firewall máy server phải cho phép Java "
        "hoặc inbound TCP 5000. Camera, microphone và speaker phụ thuộc driver hệ điều hành. STT yêu cầu thư mục model Vosk "
        "hợp lệ. Dữ liệu media lớn có thể làm tăng độ trễ vì protocol hiện chưa có chunking hoặc back-pressure.")


def chapter_3(doc):
    add_heading(doc, "CHƯƠNG 3. PHÂN TÍCH HỆ THỐNG", 1)
    add_heading(doc, "3.1. Kiến trúc tổng thể", 2)
    add_para(doc,
        "Hệ thống tổ chức theo kiến trúc client - server kết hợp phân lớp. Presentation layer chứa các JFrame và component "
        "Swing. Application layer chịu trách nhiệm socket, media và speech-to-text. Server layer quản lý kết nối, xác thực, "
        "định tuyến và persistence. Data layer dùng DAO trên một kết nối SQLite do DatabaseManager quản lý.")
    add_figure(doc, DIAGRAMS / "02-architecture.png", "Hình 3.1. Kiến trúc phân lớp và các thành phần chính")
    add_heading(doc, "3.2. Phân tích luồng xác thực", 2)
    add_para(doc,
        "Client mở một socket tạm để gửi LOGIN hoặc REGISTER. Server truy vấn username, tạo RSA key pair khi đăng ký mới "
        "hoặc so sánh password hash khi đăng nhập. Sau khi nhận kết quả thành công, socket tạm đóng; ClientUI mở socket "
        "hoạt động lâu dài và gửi CONNECT để tham gia danh sách online.")
    add_figure(doc, DIAGRAMS / "03-auth-activity.png", "Hình 3.2. Biểu đồ hoạt động đăng nhập và đăng ký", width=4.80)
    add_figure(doc, DIAGRAMS / "04-auth-sequence.png", "Hình 3.3. Biểu đồ tuần tự xác thực tài khoản")
    add_heading(doc, "3.3. Phân tích luồng tin nhắn mã hóa", 2)
    add_para(doc,
        "Khi người gửi chọn một tài khoản, client yêu cầu public key của người nhận từ server. Một AES key và IV mới được "
        "sinh cho từng tin. Ciphertext được tạo bởi AES-GCM; AES key được RSA-encrypt hai lần cho sender và receiver. "
        "Server chỉ định tuyến đối tượng Message và dự kiến lưu ciphertext cùng hai encrypted key.")
    add_figure(doc, DIAGRAMS / "05-message-sequence.png", "Hình 3.4. Biểu đồ tuần tự gửi và nhận tin nhắn mã hóa")
    add_heading(doc, "3.4. Phân tích cuộc gọi", 2)
    add_para(doc,
        "Voice và video call đều dùng signaling qua các MessageType *_CALL_REQ, *_CALL_RES và *_CALL_END. Sau khi callee "
        "chấp nhận, media handler mở thiết bị và gửi các gói AUDIO_DATA hoặc VIDEO_DATA. Server không xử lý codec mà chỉ "
        "forward gói tới receiver. UI call duy trì timer và các trạng thái mute, camera off, speaker off.")
    add_figure(doc, DIAGRAMS / "06-call-sequence.png", "Hình 3.5. Biểu đồ tuần tự voice call và video call")


def chapter_4(doc):
    add_heading(doc, "CHƯƠNG 4. THIẾT KẾ HỆ THỐNG", 1)
    add_heading(doc, "4.1. Thiết kế package và lớp", 2)
    add_para(doc,
        "Source được chia thành năm package chính. com.chat.client chứa giao diện và handler phía người dùng; "
        "com.chat.server chứa server socket và thread xử lý từng kết nối; com.chat.common định nghĩa Message, UITheme và "
        "steganography; com.chat.crypto cung cấp AES, RSA và dịch vụ crypto; com.chat.db quản lý SQLite và DAO. Cách chia "
        "này làm rõ trách nhiệm và giảm phụ thuộc chéo giữa UI với persistence.")
    add_figure(doc, DIAGRAMS / "07-class-diagram.png", "Hình 4.1. Biểu đồ lớp rút gọn và quan hệ phụ thuộc")
    class_rows = [
        ("LoginUI", "Client", "Chọn LOGIN/REGISTER, validate form, nhận crypto context."),
        ("ClientUI", "Client", "Hiển thị chat, mã hóa text, điều phối file và call."),
        ("ChatClient", "Client", "Quản lý socket, object streams và thread nhận Message."),
        ("AudioHandler", "Client", "Capture/play audio, mute/speaker, call timer."),
        ("VideoHandler", "Client", "Capture video/audio, self-preview, camera controls."),
        ("ChatServer", "Server", "Accept socket, map username -> ClientHandler, broadcast online list."),
        ("ClientHandler", "Server", "Xác thực, dispatch protocol, forward và persistence."),
        ("DatabaseManager", "DB", "Singleton connection và tạo schema SQLite."),
        ("UserDAO / MessageDAO", "DB", "CRUD account và ciphertext."),
        ("AESUtil / RSAUtil", "Crypto", "Mã hóa lai, key serialization và Base64."),
    ]
    add_table(doc, ["Lớp", "Nhóm", "Trách nhiệm chính"], class_rows,
              [1900, 1200, 5972], "Bảng 4.1. Vai trò của các lớp trọng tâm")

    add_heading(doc, "4.2. Thiết kế protocol Message", 2)
    add_para(doc,
        "Mọi dữ liệu nghiệp vụ được đóng gói trong lớp Message implements Serializable. Trường type chọn nhánh xử lý; "
        "sender và receiver định tuyến; content chứa payload thông thường; fileData chứa binary; nhóm trường encryption "
        "chứa ciphertext, IV, userId và hai encrypted AES key. serialVersionUID được tăng khi protocol xác thực thay đổi.")
    protocol_rows = [
        ("Xác thực", "LOGIN, LOGIN_OK, LOGIN_FAIL, REGISTER, REGISTER_OK, REGISTER_FAIL"),
        ("Kết nối", "CONNECT, DISCONNECT, USER_LIST"),
        ("Nội dung", "TEXT, ENCRYPTED_TEXT, ICON, FILE, STEGANOGRAPHY"),
        ("Trao đổi khóa", "KEY_REQUEST, KEY_RESPONSE"),
        ("Voice", "VOICE_CALL_REQ, VOICE_CALL_RES, VOICE_CALL_END, AUDIO_DATA"),
        ("Video", "VIDEO_CALL_REQ, VIDEO_CALL_RES, VIDEO_CALL_END, VIDEO_DATA"),
    ]
    add_table(doc, ["Nhóm", "MessageType"], protocol_rows, [1800, 7272], "Bảng 4.2. Phân nhóm protocol")

    add_heading(doc, "4.3. Thiết kế cơ sở dữ liệu", 2)
    add_para(doc,
        "SQLite được đặt tại user.home/chat_server.db để server và client không nhầm file khi chạy từ các working directory "
        "khác nhau. Schema gồm users, messages và message_keys. users lưu định danh, password hash và RSA key pair. "
        "messages lưu ciphertext cùng IV. message_keys lưu encrypted AES key theo từng user được phép đọc một message.")
    add_figure(doc, DIAGRAMS / "08-erd.png", "Hình 4.2. Biểu đồ ERD của cơ sở dữ liệu SQLite")
    db_rows = [
        ("users", "id", "username, password_hash, public_key, private_key, created_at", "Tài khoản và key pair"),
        ("messages", "id", "sender_id, receiver_id, encrypted_content, iv, created_at", "Ciphertext từng tin"),
        ("message_keys", "id", "message_id, user_id, encrypted_aes_key", "Khóa AES mã hóa theo user"),
    ]
    add_table(doc, ["Bảng", "Khóa chính", "Thuộc tính", "Ý nghĩa"], db_rows,
              [1300, 1200, 4500, 2072], "Bảng 4.3. Mô tả lược đồ dữ liệu")
    add_callout(doc, "Điểm cần hoàn thiện",
        "Client hiện đặt receiverId = -1 khi gửi ENCRYPTED_TEXT trong khi ClientHandler chưa ánh xạ username thành id trước "
        "khi INSERT. Vì foreign key được bật, persistence lịch sử có thể thất bại dù tin vẫn được forward thời gian thực.",
        fill="FEF3F2", accent=RED)

    add_heading(doc, "4.4. Thiết kế bảo mật", 2)
    for item in [
        "Mật khẩu: client băm SHA-256 trước khi truyền; server so sánh password_hash trong bảng users.",
        "Nội dung: AES/GCM/NoPadding, key 256 bit, IV 12 byte và authentication tag 128 bit.",
        "Khóa nội dung: RSA 2048 bit với OAEP SHA-256 và MGF1 SHA-256.",
        "Khả năng đọc lại: cùng AES key được mã hóa bằng public key sender và receiver.",
        "Tính toàn vẹn: GCM phát hiện ciphertext bị sửa thông qua AEAD tag.",
    ]:
        add_bullet(doc, item)
    add_para(doc,
        "Mô hình hiện tại mang tính minh họa học thuật, chưa đạt zero-knowledge end-to-end encryption vì server lưu private key "
        "và truyền key trên socket chưa có TLS. Phiên bản production cần Argon2id hoặc BCrypt cho password, TLS cho transport, "
        "private key được mã hóa bằng khóa dẫn xuất từ password và cơ chế xác minh danh tính thiết bị.")

    add_heading(doc, "4.5. Thiết kế trạng thái cuộc gọi", 2)
    add_figure(doc, DIAGRAMS / "10-call-state.png", "Hình 4.3. Biểu đồ trạng thái cuộc gọi")


def chapter_5(doc):
    add_heading(doc, "CHƯƠNG 5. CÀI ĐẶT VÀ XÂY DỰNG CHỨC NĂNG", 1)
    add_heading(doc, "5.1. Công nghệ sử dụng", 2)
    tech_rows = [
        ("Java", "17 target / JDK 21 kiểm thử", "Ngôn ngữ, socket, thread, Swing, Sound API"),
        ("Maven", "pom.xml", "Quản lý dependency và build"),
        ("FlatLaf", "3.4", "Look and Feel hiện đại"),
        ("flatlaf-extras", "3.4", "Render SVG vector"),
        ("SQLite JDBC", "3.45.3.0", "Lưu account và ciphertext"),
        ("Webcam Capture", "0.3.12", "Lấy frame webcam"),
        ("Vosk", "0.3.32", "Speech-to-Text offline"),
        ("JNA", "5.13.0", "Hỗ trợ native cho Vosk"),
    ]
    add_table(doc, ["Công nghệ", "Phiên bản", "Vai trò"], tech_rows,
              [1900, 1800, 5372], "Bảng 5.1. Technology stack")

    add_heading(doc, "5.2. Giao diện đăng nhập và đăng ký", 2)
    add_para(doc,
        "Màn hình authentication sử dụng hai chế độ hiển thị trong cùng JFrame. Sign in yêu cầu server address, port, username "
        "và password. Create account hiển thị thêm confirm password, kiểm tra username tối thiểu ba ký tự và password tối "
        "thiểu sáu ký tự. Tác vụ socket chạy trong SwingWorker để không khóa UI.")
    add_figure(doc, SCREENS / "register.png", "Hình 5.1. Giao diện Create account", width=5.85)

    add_heading(doc, "5.3. Giao diện chat", 2)
    add_para(doc,
        "ClientUI gồm sidebar tài khoản hiện tại, tìm kiếm hội thoại và danh sách online; header hiển thị người đang chat cùng "
        "nút voice/video; vùng giữa chứa message bubble; composer có icon SVG cho sticker, ảnh, file, steganography và STT. "
        "Mỗi người dùng có một chat panel riêng được cache trong Map để chuyển hội thoại nhanh.")
    add_figure(doc, SCREENS / "chat.png", "Hình 5.2. Giao diện chat với icon SVG và bubble mã hóa", width=5.85)

    add_heading(doc, "5.4. Gửi tin nhắn, ảnh và tệp", 2)
    add_para(doc,
        "sendText() cập nhật bubble lạc quan, sau đó SwingWorker lấy public key, tạo AES key/IV, mã hóa và gửi Message. "
        "Ảnh được đọc thành byte array, nhận và hiển thị bằng ImageIcon. Tệp được chuyển nguyên byte và người nhận chọn vị trí "
        "lưu. ICON được dùng cho sticker và ảnh; FILE dành cho tài liệu; STEGANOGRAPHY chứa ảnh PNG đã nhúng dữ liệu.")
    add_heading(doc, "5.5. Steganography và Speech-to-Text", 2)
    add_para(doc,
        "SteganoUtils thay bit thấp nhất của kênh RGB để lưu độ dài và chuỗi password::message. Đây là kỹ thuật giấu dữ liệu, "
        "không phải thuật toán mã hóa mạnh; password chỉ đóng vai trò kiểm tra. SttHandler mở microphone, chọn sample rate phù "
        "hợp và đưa audio vào Vosk Recognizer. Văn bản cuối được gửi dưới dạng TEXT với tiền tố Transcribed.")

    add_heading(doc, "5.6. Voice call và video call", 2)
    add_para(doc,
        "AudioHandler capture PCM mono 8 kHz, 16 bit và gửi AUDIO_DATA. UI voice call có timer, mute, speaker và end. "
        "VideoHandler capture webcam, scale frame 480 x 360, nén JPEG và gửi khoảng 10 FPS; audio được capture song song. "
        "UI video có remote stage, self-preview, mic, camera, speaker và end. Tất cả control sử dụng SVG vector.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    voice_shape = r.add_picture(str(SCREENS / "voice-call.png"), width=Inches(2.35))
    voice_shape._inline.docPr.set("descr", "Giao diện voice call")
    voice_shape._inline.docPr.set("title", "Hình 5.3. Giao diện voice call")
    video_shape = r.add_picture(str(SCREENS / "video-call.png"), width=Inches(2.75))
    video_shape._inline.docPr.set("descr", "Giao diện video call")
    video_shape._inline.docPr.set("title", "Hình 5.4. Giao diện video call")
    p.paragraph_format.keep_with_next = True
    doc.add_paragraph("Hình 5.3 và 5.4. Giao diện voice call và video call", style="Caption")


def chapter_6(doc):
    add_heading(doc, "CHƯƠNG 6. KIỂM THỬ VÀ ĐÁNH GIÁ", 1)
    add_heading(doc, "6.1. Môi trường kiểm thử", 2)
    add_para(doc,
        "Dự án được kiểm tra trên Windows, Java 21 runtime và Maven tích hợp IntelliJ; compiler target là Java 17. Build chạy "
        "lệnh Maven test, biên dịch 20 source file và copy 17 SVG resource. Các UI đăng ký, chat, voice call và video call "
        "được render trực tiếp từ source để kiểm tra font, icon, khoảng cách và clipping.")
    add_heading(doc, "6.2. Ca kiểm thử", 2)
    tests = [
        ("TC-01", "Build Maven", "mvn test", "BUILD SUCCESS", "Đạt"),
        ("TC-02", "Đăng ký mới", "Username chưa tồn tại", "REGISTER_OK", "Đạt"),
        ("TC-03", "Đăng ký trùng", "Username đã tồn tại", "REGISTER_FAIL", "Đạt"),
        ("TC-04", "Đăng nhập đúng", "Đúng password", "LOGIN_OK", "Đạt"),
        ("TC-05", "Đăng nhập sai", "Sai password", "LOGIN_FAIL", "Đạt"),
        ("TC-06", "Account không có", "Username mới ở Sign in", "LOGIN_FAIL", "Đạt"),
        ("TC-07", "Danh sách online", "Hai client CONNECT", "USER_LIST cập nhật", "Đạt thủ công"),
        ("TC-08", "Chat text", "Hai user online", "Forward và decrypt", "Đạt thủ công"),
        ("TC-09", "Persistence", "ENCRYPTED_TEXT", "Cần resolve receiverId", "Chưa hoàn thiện"),
        ("TC-10", "Gửi ảnh / file", "File hợp lệ", "Nhận và lưu", "Đạt thủ công"),
        ("TC-11", "Voice call", "Có microphone/speaker", "Audio hai chiều", "Cần thiết bị"),
        ("TC-12", "Video call", "Có webcam", "Video + audio", "Cần thiết bị"),
        ("TC-13", "STT", "Có Vosk model", "Trả văn bản", "Cần thiết bị"),
        ("TC-14", "LAN", "Hai máy cùng Wi-Fi", "Kết nối IPv4:5000", "Cần môi trường"),
    ]
    add_table(doc, ["Mã", "Nội dung", "Điều kiện", "Kỳ vọng", "Kết quả"], tests,
              [800, 1900, 2250, 2350, 1772], "Bảng 6.1. Ma trận kiểm thử")
    add_heading(doc, "6.3. Đánh giá ưu điểm", 2)
    for item in [
        "Kiến trúc package rõ ràng, thể hiện đầy đủ client, server, DAO, crypto và media.",
        "UI hiện đại, icon SVG nhất quán, hiển thị tiếng Việt và thao tác call trực quan.",
        "Tích hợp nhiều chủ đề học thuật trong một sản phẩm: socket, thread, DB, crypto, audio/video và STT.",
        "Account flow rõ ràng và đã có smoke test protocol độc lập.",
        "Server có console, client có tìm kiếm online user và cache panel hội thoại.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "6.4. Hạn chế và rủi ro", 2)
    risks = [
        ("R-01", "receiverId = -1 làm persistence message thất bại", "Cao", "Server lookup UserRecord theo receiver username trước INSERT."),
        ("R-02", "UI chưa tải lịch sử chat từ DB", "Cao", "Bổ sung HISTORY_REQUEST/HISTORY_RESPONSE và decrypt client-side."),
        ("R-03", "Private key lưu rõ tại server", "Cao", "Mã hóa private key bằng khóa dẫn xuất từ password; không gửi qua socket rõ."),
        ("R-04", "Java ObjectInputStream với input không tin cậy", "Cao", "Chuyển sang JSON/Protobuf và whitelist schema."),
        ("R-05", "Transport chưa TLS", "Cao", "Dùng SSLSocket hoặc TLS termination."),
        ("R-06", "Không giới hạn file/frame", "Trung bình", "Giới hạn kích thước, chunking, timeout và back-pressure."),
        ("R-07", "ObjectOutputStream có thể bị gọi đồng thời", "Trung bình", "Serialize write qua queue hoặc synchronized sendMessage."),
    ]
    add_table(doc, ["Mã", "Rủi ro", "Mức", "Biện pháp"], risks,
              [750, 3000, 1000, 4322], "Bảng 6.2. Danh sách rủi ro kỹ thuật")


def chapter_7(doc):
    add_heading(doc, "CHƯƠNG 7. TRIỂN KHAI VÀ HƯỚNG DẪN SỬ DỤNG", 1)
    add_heading(doc, "7.1. Mô hình triển khai", 2)
    add_para(doc,
        "Server chỉ cần chạy trên một máy. Máy server cũng có thể chạy một client cục bộ bằng localhost. Các máy khác trong "
        "cùng Wi-Fi chỉ cần chạy LoginUI và nhập IPv4 của máy server. ServerSocket(5000) lắng nghe trên các interface mạng "
        "khả dụng, do đó có thể phục vụ localhost và LAN đồng thời.")
    add_figure(doc, DIAGRAMS / "09-deployment.png", "Hình 7.1. Biểu đồ triển khai hai máy trong mạng LAN")
    add_heading(doc, "7.2. Yêu cầu môi trường", 2)
    for item in [
        "JDK 17 trở lên; khuyến nghị JDK 21 LTS.",
        "Maven hoặc Maven tích hợp IntelliJ IDEA.",
        "Quyền truy cập microphone, speaker và webcam nếu dùng call/STT.",
        "Model Vosk trong thư mục model nếu dùng Speech-to-Text.",
        "Firewall cho phép Java hoặc inbound TCP port 5000 trên máy server.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "7.3. Build và chạy", 2)
    add_code_block(doc,
        "# Build\n"
        "mvn clean test\n\n"
        "# Server tự khởi động cổng 5000\n"
        "mvn exec:java -Dexec.mainClass=com.chat.server.ServerUI -Dexec.args=--start\n\n"
        "# Client\n"
        "mvn exec:java -Dexec.mainClass=com.chat.client.LoginUI")
    add_heading(doc, "7.4. Quy trình sử dụng", 2)
    steps = [
        "Mở ServerUI và bảo đảm trạng thái ONLINE, port 5000.",
        "Tại client, chọn Create account cho lần sử dụng đầu tiên; nhập server IP, port, username và password.",
        "Lần sau chọn Sign in với thông tin đã đăng ký.",
        "Chọn tài khoản online trong sidebar, nhập nội dung và bấm Send.",
        "Dùng các icon để gửi sticker, ảnh, file, ảnh giấu tin hoặc Speech-to-Text.",
        "Dùng nút phone/video ở header; bên nhận chấp nhận yêu cầu và điều khiển mic/camera/speaker.",
    ]
    for item in steps:
        add_number(doc, item)
    add_heading(doc, "7.5. Xử lý lỗi thường gặp", 2)
    troubleshooting = [
        ("Connection refused", "Server chưa chạy, sai IP/port hoặc firewall chặn."),
        ("Account does not exist", "Đang ở Sign in nhưng chưa đăng ký; chuyển Create account."),
        ("Incorrect password", "Mật khẩu không khớp password_hash đã lưu."),
        ("No webcam detected", "Kiểm tra webcam, quyền riêng tư Windows và ứng dụng đang chiếm camera."),
        ("Microphone unavailable", "Kiểm tra input device và định dạng audio được hỗ trợ."),
        ("Vosk model not found", "Giải nén model đúng vào thư mục model ở project root."),
    ]
    add_table(doc, ["Hiện tượng", "Cách xử lý"], troubleshooting, [2800, 6272], "Bảng 7.1. Troubleshooting")


def chapter_8(doc):
    add_heading(doc, "CHƯƠNG 8. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", 1)
    add_heading(doc, "8.1. Kết luận", 2)
    add_para(doc,
        "Đồ án đã xây dựng được một ứng dụng chat desktop có kiến trúc hoàn chỉnh và nhiều chức năng tích hợp. Qua quá trình "
        "thực hiện, các nội dung socket TCP, xử lý đa luồng, serialization, Swing event model, SQLite DAO, mã hóa lai, "
        "capture/play media và speech recognition được kết nối trong cùng một sản phẩm. Giao diện cuối cùng có khả năng "
        "đăng ký, đăng nhập, chat, gửi tệp, gọi thoại và gọi video với trải nghiệm nhất quán.")
    add_para(doc,
        "Giá trị lớn nhất của dự án không chỉ nằm ở số lượng tính năng mà còn ở việc làm rõ ranh giới giữa client và server: "
        "client chịu trách nhiệm trải nghiệm và xử lý nội dung nhạy cảm; server duy trì kết nối, định tuyến và persistence. "
        "Các hạn chế được ghi nhận trung thực tạo cơ sở cho vòng phát triển tiếp theo thay vì che giấu rủi ro kỹ thuật.")
    add_heading(doc, "8.2. Hướng phát triển", 2)
    roadmap = [
        ("Giai đoạn 1", "Hoàn thiện receiverId, history API, load/decrypt hội thoại và trạng thái đã đọc."),
        ("Giai đoạn 2", "TLS, BCrypt/Argon2id, mã hóa private key, session token và chống giả mạo CONNECT."),
        ("Giai đoạn 3", "Protocol JSON/Protobuf, message queue, chunked upload, giới hạn dữ liệu và retry."),
        ("Giai đoạn 4", "Group chat, contact, offline message, notification, search nội dung cục bộ."),
        ("Giai đoạn 5", "TURN/STUN hoặc WebRTC cho media, codec tối ưu và adaptive bitrate."),
        ("Giai đoạn 6", "Test tự động, CI/CD, đóng gói jpackage và installer cho Windows."),
    ]
    add_table(doc, ["Giai đoạn", "Nội dung"], roadmap, [1800, 7272], "Bảng 8.1. Roadmap đề xuất")
    add_callout(doc, "Định hướng ưu tiên",
        "Ưu tiên sửa persistence và security boundary trước khi bổ sung thêm chức năng. Một hệ thống ít tính năng nhưng bảo "
        "toàn dữ liệu đúng sẽ có giá trị hơn một hệ thống nhiều tính năng nhưng không khôi phục được lịch sử hoặc bảo vệ key.",
        fill="ECFDF3", accent=TEAL)


def references(doc):
    add_heading(doc, "TÀI LIỆU THAM KHẢO", 1)
    refs = [
        "Oracle. Java Platform, Standard Edition Documentation. https://docs.oracle.com/en/java/",
        "SQLite Documentation. https://www.sqlite.org/docs.html",
        "JFormDesigner. FlatLaf Documentation. https://www.formdev.com/flatlaf/",
        "Alpha Cephei. Vosk Speech Recognition Toolkit. https://alphacephei.com/vosk/",
        "Sarxos. Webcam Capture API. https://github.com/sarxos/webcam-capture",
        "NIST Special Publication 800-38D: Galois/Counter Mode (GCM).",
        "RFC 8017: PKCS #1 - RSA Cryptography Specifications Version 2.2.",
        "Tài liệu source ChatClientJava, commit 1ded34a, nhánh main.",
    ]
    for index, item in enumerate(refs, start=1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.28)
        p.paragraph_format.space_after = Pt(4)
        set_run_font(p.add_run(f"[{index}] {item}"), size=13, color=INK)


def appendices(doc):
    add_heading(doc, "PHỤ LỤC", 1)
    add_heading(doc, "Phụ lục A. Cấu trúc source", 2)
    add_code_block(doc,
        "src/main/java/com/chat/\n"
        "├── client/\n"
        "│   ├── LoginUI, ClientUI, ChatClient\n"
        "│   └── AudioHandler, VideoHandler, SttHandler\n"
        "├── server/   ServerUI, ChatServer, ClientHandler\n"
        "├── common/   Message, UITheme, SteganoUtils\n"
        "├── crypto/   AESUtil, RSAUtil, KeyManager\n"
        "│             CryptoMessageService\n"
        "└── db/       DatabaseManager, UserDAO, MessageDAO\n"
        "src/main/resources/icons/   SVG icons\n"
        "model/                       Vosk model")
    add_heading(doc, "Phụ lục B. DDL SQLite", 2)
    add_code_block(doc,
        "CREATE TABLE users (\n"
        "  id INTEGER PRIMARY KEY AUTOINCREMENT,\n"
        "  username TEXT NOT NULL UNIQUE,\n"
        "  password_hash TEXT NOT NULL,\n"
        "  public_key TEXT NOT NULL, private_key TEXT NOT NULL,\n"
        "  created_at DATETIME DEFAULT CURRENT_TIMESTAMP);\n\n"
        "CREATE TABLE messages (\n"
        "  id INTEGER PRIMARY KEY AUTOINCREMENT,\n"
        "  sender_id INTEGER NOT NULL, receiver_id INTEGER NOT NULL,\n"
        "  encrypted_content TEXT NOT NULL, iv TEXT NOT NULL,\n"
        "  created_at DATETIME DEFAULT CURRENT_TIMESTAMP,\n"
        "  FOREIGN KEY(sender_id) REFERENCES users(id),\n"
        "  FOREIGN KEY(receiver_id) REFERENCES users(id));\n\n"
        "CREATE TABLE message_keys (\n"
        "  id INTEGER PRIMARY KEY AUTOINCREMENT,\n"
        "  message_id INTEGER NOT NULL, user_id INTEGER NOT NULL,\n"
        "  encrypted_aes_key TEXT NOT NULL,\n"
        "  FOREIGN KEY(message_id) REFERENCES messages(id),\n"
        "  FOREIGN KEY(user_id) REFERENCES users(id));")
    add_heading(doc, "Phụ lục C. Tham số cấu hình chính", 2)
    config_rows = [
        ("Server port", "5000", "ServerUI / LoginUI"),
        ("DB path", "${user.home}/chat_server.db", "DatabaseManager"),
        ("AES", "256-bit, GCM, IV 12 byte", "AESUtil"),
        ("RSA", "2048-bit, OAEP SHA-256", "RSAUtil"),
        ("Audio", "8 kHz, 16 bit, mono", "AudioHandler / VideoHandler"),
        ("Video", "480 x 360, JPEG, khoảng 10 FPS", "VideoHandler"),
        ("Java target", "17", "pom.xml"),
    ]
    add_table(doc, ["Tham số", "Giá trị", "Vị trí"], config_rows,
              [1900, 3800, 3372], "Bảng C.1. Cấu hình kỹ thuật")
    add_heading(doc, "Phụ lục D. Checklist nghiệm thu", 2)
    for item in [
        "Build Maven không lỗi và resource SVG được copy.",
        "Server khởi động, SQLite schema được tạo, port 5000 lắng nghe.",
        "Đăng ký mới, đăng ký trùng, đăng nhập đúng và sai đều trả đúng protocol.",
        "Hai client hiển thị nhau trong online list và chat được hai chiều.",
        "File, ảnh, steganography, voice, video và STT được kiểm tra theo thiết bị thực tế.",
    ]:
        add_bullet(doc, item)


def build_report(output_path, toc_pages=None):
    generate_all_diagrams()
    toc_pages = toc_pages or {}
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_toc(doc, toc_pages)
    add_abstract(doc)
    page_break(doc)
    add_lists(doc)
    chapter_1(doc)
    chapter_2(doc)
    chapter_3(doc)
    chapter_4(doc)
    chapter_5(doc)
    chapter_6(doc)
    chapter_7(doc)
    chapter_8(doc)
    references(doc)
    appendices(doc)

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--out", default=str(DOCS / "Bao_cao_do_an_ChatClientJava.docx"))
    parser.add_argument("--toc-pages", default=None)
    args = parser.parse_args()
    toc = {}
    if args.toc_pages and Path(args.toc_pages).exists():
        toc = json.loads(Path(args.toc_pages).read_text(encoding="utf-8"))
    build_report(Path(args.out), toc)
    print(args.out)


if __name__ == "__main__":
    main()
